"""
DOCX exporter for Word documents.
"""
from typing import Dict
import logging
from docx import Document
from docx.shared import Pt
from app.exporters.base import BaseExporter

logger = logging.getLogger(__name__)


class DOCXExporter(BaseExporter):
    """Export OCR results to DOCX format."""

    def export(self, data: Dict, output_path: str) -> str:
        """
        Export to DOCX file.

        Args:
            data: OCR result data
            output_path: Output file path

        Returns:
            Path to created file
        """
        self._ensure_output_dir(output_path)

        try:
            doc = Document()

            # Add text content
            text_data = data.get("text", {})
            full_text = text_data.get("full_text", "")

            if full_text:
                paragraphs = full_text.split("\n\n")
                for para_text in paragraphs:
                    if para_text.strip():
                        p = doc.add_paragraph(para_text.strip())
                        p.style.font.size = Pt(11)

            # Add tables
            tables = data.get("tables", [])
            if tables:
                doc.add_page_break()
                doc.add_heading("Таблицы", level=1)

                for idx, table_data in enumerate(tables, 1):
                    doc.add_heading(f"Таблица {idx}", level=2)

                    cells = table_data.get("cells", [])
                    if cells:
                        # Create table
                        rows = len(cells)
                        cols = max(len(row) for row in cells) if cells else 0

                        table = doc.add_table(rows=rows, cols=cols)
                        table.style = 'Light Grid Accent 1'

                        # Fill table
                        for i, row_data in enumerate(cells):
                            row = table.rows[i]
                            for j, cell_value in enumerate(row_data):
                                if j < len(row.cells):
                                    row.cells[j].text = str(cell_value)

                        # Apply merged cells
                        merged_regions = table_data.get("merged_regions", [])
                        for merge in merged_regions:
                            try:
                                start_row = merge['row']
                                start_col = merge['col']
                                rowspan = merge['rowspan']
                                colspan = merge['colspan']

                                # Only merge if it's actually a merged region
                                if rowspan > 1 or colspan > 1:
                                    # Get start and end cells
                                    start_cell = table.rows[start_row].cells[start_col]
                                    end_row = min(start_row + rowspan - 1, rows - 1)
                                    end_col = min(start_col + colspan - 1, cols - 1)
                                    end_cell = table.rows[end_row].cells[end_col]

                                    # Merge cells
                                    start_cell.merge(end_cell)
                            except Exception as e:
                                logger.warning(f"Could not merge cells in DOCX: {e}")

                        doc.add_paragraph()  # Add spacing

            doc.save(output_path)
            logger.info(f"Exported to DOCX: {output_path}")
            return output_path

        except Exception as e:
            logger.error(f"Error exporting to DOCX: {e}")
            raise
