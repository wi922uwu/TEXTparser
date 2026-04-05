"""Export text to TXT and DOCX formats."""
from pathlib import Path
from docx import Document as DocxDocument
from docx.shared import Pt

from backend.models.schemas import TextRegion, TableRegion


def export_txt(
    text_regions: list[TextRegion],
    table_regions: list[TableRegion],
    output_path: Path,
) -> Path:
    """Export extracted content as plain text."""
    lines = []

    for region in text_regions:
        lines.append(region.text)

    if table_regions and text_regions:
        lines.append("")
        lines.append("--- TABLES ---")
        lines.append("")

    for table in table_regions:
        for row in table.cells:
            lines.append(" | ".join(str(cell) for cell in row))
        lines.append("")

    output_path.write_text("\n".join(lines), encoding="utf-8")
    return output_path


def export_docx(
    text_regions: list[TextRegion],
    table_regions: list[TableRegion],
    output_path: Path,
) -> Path:
    """Export extracted content as DOCX with formatted tables."""
    doc = DocxDocument()

    for region in text_regions:
        para = doc.add_paragraph(region.text)
        for run in para.runs:
            run.font.size = Pt(11)

    for table_data in table_regions:
        if not table_data.cells:
            continue

        table = doc.add_table(
            rows=len(table_data.cells),
            cols=max(len(row) for row in table_data.cells),
        )
        table.style = "Table Grid"

        for i, row_cells in enumerate(table_data.cells):
            for j, cell_text in enumerate(row_cells):
                table.cell(i, j).text = str(cell_text)
                for paragraph in table.cell(i, j).paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(10)

        doc.add_paragraph()

    doc.save(str(output_path))
    return output_path
